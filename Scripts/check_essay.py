import sys
import argparse
import json
import os
from datetime import datetime
from docx import Document
from docx.shared import Pt
import tkinter as tk
from tkinter import filedialog
from collections import defaultdict
from docx.oxml.ns import qn


def select_files():
    root = tk.Tk()
    root.withdraw()
    file_paths = filedialog.askopenfilenames(
        title="Виберіть файли рефератів",
        filetypes=[("Word документи", "*.docx"), ("Всі файли", "*.*")]
    )
    return list(file_paths)


def get_page_numbers(doc):
    page_numbers = []
    current_page = 1
    for paragraph in doc.paragraphs:
        has_page_break = False
        for run in paragraph.runs:
            brs = run._element.findall(qn('w:br'))
            for br in brs:
                if br.get(qn('w:type')) == 'page':
                    has_page_break = True
                    break
            if has_page_break:
                break
        if has_page_break:
            current_page += 1
        page_numbers.append(current_page)
    return page_numbers


def check_essay(file_path, expected_font='Times New Roman', expected_size=14):
    errors_by_paragraph = defaultdict(lambda: {'font_issues': [], 'size_issues': [], 'text': '', 'page': 1})

    try:
        doc = Document(file_path)
    except FileNotFoundError:
        return [f" Файл не знайдено: {file_path}"]
    except Exception as e:
        return [f" Помилка відкриття файлу: {e}"]

    page_numbers = get_page_numbers(doc)

    for para_index, paragraph in enumerate(doc.paragraphs):
        para_text = paragraph.text
        if not para_text.strip():
            continue

        errors_by_paragraph[para_index + 1]['text'] = para_text[:200] + ('...' if len(para_text) > 200 else '')
        errors_by_paragraph[para_index + 1]['page'] = page_numbers[para_index]

        runs = list(paragraph.runs)
        for i, run in enumerate(runs):
            font_name = run.font.name
            font_size = run.font.size
            fragment_text = run.text

            start_pos = 0
            for j in range(i):
                start_pos += len(runs[j].text)
            end_pos = start_pos + len(fragment_text)

            is_space = fragment_text and not fragment_text.strip()

            if is_space:
                prev_text = ''
                next_text = ''
                for j in range(i - 1, -1, -1):
                    if runs[j].text.strip():
                        prev_text = runs[j].text.strip()
                        break
                for j in range(i + 1, len(runs)):
                    if runs[j].text.strip():
                        next_text = runs[j].text.strip()
                        break
                if prev_text or next_text:
                    display_fragment = f"пробіл між \"{prev_text}\" та \"{next_text}\""
                else:
                    display_fragment = "<пробіл без контексту>"
            else:
                display_fragment = fragment_text[:70] + ('...' if len(fragment_text) > 70 else '')

            if font_name and font_name.lower() != expected_font.lower():
                errors_by_paragraph[para_index + 1]['font_issues'].append({
                    'actual': font_name,
                    'fragment': display_fragment,
                    'text': fragment_text,
                    'start': start_pos,
                    'end': end_pos
                })

            if font_size and font_size != Pt(expected_size):
                errors_by_paragraph[para_index + 1]['size_issues'].append({
                    'actual': font_size.pt,
                    'fragment': display_fragment,
                    'text': fragment_text,
                    'start': start_pos,
                    'end': end_pos
                })

    json_errors = []
    error_id = 1
    for para_num, issues in errors_by_paragraph.items():
        if issues['font_issues'] or issues['size_issues']:
            for issue in issues['font_issues']:
                json_errors.append({
                    'id': f"err_{error_id:03d}",
                    'category': 'font',
                    'severity': 'error',
                    'title': f"Невірний шрифт: {issue['actual']}",
                    'paragraph_index': para_num,
                    'page': issues['page'],
                    'paragraph_text': issues['text'],
                    'highlight': {
                        'start': issue['start'],
                        'end': issue['end']
                    },
                    'found': issue['actual'],
                    'expected': expected_font,
                    'suggestions': [expected_font]
                })
                error_id += 1

            for issue in issues['size_issues']:
                json_errors.append({
                    'id': f"err_{error_id:03d}",
                    'category': 'font_size',
                    'severity': 'error',
                    'title': f"Невірний розмір шрифту: {issue['actual']}pt",
                    'paragraph_index': para_num,
                    'page': issues['page'],
                    'paragraph_text': issues['text'],
                    'highlight': {
                        'start': issue['start'],
                        'end': issue['end']
                    },
                    'found': f"{issue['actual']}pt",
                    'expected': f"{expected_size}pt",
                    'suggestions': [f"{expected_size}pt"]
                })
                error_id += 1

    return json_errors


def generate_json_report(file_path, errors, font, size):
    filename = os.path.basename(file_path)

    by_severity = {'error': 0, 'warning': 0}
    by_category = {}

    for err in errors:
        by_severity[err['severity']] = by_severity.get(err['severity'], 0) + 1
        category = err['category']
        by_category[category] = by_category.get(category, 0) + 1

    report = {
        "file_name": filename,
        "analyzed_at": datetime.now().isoformat(),
        "summary": {
            "total_errors": len(errors),
            "by_severity": by_severity,
            "by_category": by_category
        },
        "errors": errors
    }
    return report


def save_json_report(report, base_filename):
    reports_dir = "../reports"
    os.makedirs(reports_dir, exist_ok=True)

    safe_name = os.path.splitext(os.path.basename(base_filename))[0]
    safe_name = "".join(c if c.isalnum() or c in "._-" else "_" for c in safe_name)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_filename = f"{safe_name}_{timestamp}.json"
    output_path = os.path.join(reports_dir, output_filename)

    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(report, f, ensure_ascii=False, indent=2)

    return output_path


def print_report(file_path, errors, font, size):
    print(f"\n **Файл: {file_path}**")
    if not errors:
        print(" Всі текстові елементи мають шрифт Times New Roman 14.")
    else:
        print(f" Знайдено {len(errors)} помилок оформлення:")
        for err in errors:
            page_info = f", сторінка {err['page']}" if err['page'] else ""
            print(f"      Абзац {err['paragraph_index']}{page_info} — {err['category']}: {err['title']}")
            print(f"         у фрагменті: {err['paragraph_text'][:50]}...")
            print(f"         (очікувалося {err['expected']}, отримано {err['found']})\n")


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description='Перевірка оформлення рефератів (шрифт Times New Roman 14)')
    parser.add_argument('files', nargs='*',
                        help="Шляхи до файлів .docx (можна вказати декілька). Якщо не вказано, відкриється діалог вибору.")
    parser.add_argument('--font', default='Times New Roman', help='Очікуваний шрифт')
    parser.add_argument('--size', type=int, default=14, help='Очікуваний розмір шрифту')
    parser.add_argument('--no-json', action='store_true', help='Не зберігати JSON-звіт (тільки текстовий вивід)')

    args = parser.parse_args()

    files_to_check = []

    if not args.files:
        files_to_check = select_files()
        if not files_to_check:
            print("Файли не вибрано. Вихід.")
            sys.exit(1)
    else:
        files_to_check = args.files

    for file_path in files_to_check:
        errors = check_essay(file_path, args.font, args.size)
        print_report(file_path, errors, args.font, args.size)

        if not args.no_json:
            report = generate_json_report(file_path, errors, args.font, args.size)
            json_path = save_json_report(report, file_path)
            print(f" JSON-звіт збережено: {json_path}")